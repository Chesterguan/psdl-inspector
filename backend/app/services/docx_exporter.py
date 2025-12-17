"""Word document generator for IRB preparation with AI enrichment."""

from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE


def generate_irb_document(
    scenario_info: Dict[str, Any],
    governance_data: Dict[str, Optional[str]],
    enriched_content: Optional[Dict[str, str]] = None,
) -> bytes:
    """Generate Word document for IRB preparation.

    Args:
        scenario_info: Auto-derived info from PSDL scenario
        governance_data: User-provided governance narrative
        enriched_content: AI-generated enriched content (optional)

    Returns:
        bytes: Word document content
    """
    doc = Document()

    # Set up styles
    _setup_styles(doc)

    # Title Page
    _add_title_page(doc, scenario_info)

    # Table of Contents placeholder
    doc.add_paragraph()
    toc_heading = doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph('1. Executive Summary')
    doc.add_paragraph('2. Clinical Background')
    doc.add_paragraph('3. Algorithm Overview')
    doc.add_paragraph('4. Data Elements')
    doc.add_paragraph('5. Algorithm Logic')
    doc.add_paragraph('6. Clinical Workflow Integration')
    doc.add_paragraph('7. Safety Considerations')
    doc.add_paragraph('8. Limitations')
    doc.add_paragraph('9. Recommendation')
    doc.add_paragraph('Appendix A: Technical Specification')
    doc.add_page_break()

    # Section 1: Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    if enriched_content and enriched_content.get('executive_summary'):
        _add_narrative_paragraphs(doc, enriched_content['executive_summary'])
    else:
        # Fallback to basic summary
        clinical_summary = governance_data.get('clinical_summary')
        if clinical_summary:
            doc.add_paragraph(clinical_summary)
        else:
            _add_placeholder(doc, 'Executive summary will be generated based on provided clinical context.')
    doc.add_paragraph()

    # Section 2: Clinical Background
    doc.add_heading('2. Clinical Background', level=1)
    if enriched_content and enriched_content.get('clinical_background'):
        _add_narrative_paragraphs(doc, enriched_content['clinical_background'])
    else:
        desc = scenario_info.get('description')
        if desc:
            doc.add_paragraph(desc)
        _add_placeholder(doc, 'Clinical background explaining the condition being monitored.')
    doc.add_paragraph()

    # Section 3: Algorithm Overview
    doc.add_heading('3. Algorithm Overview', level=1)

    # Basic info table
    overview_table = doc.add_table(rows=4, cols=2)
    overview_table.style = 'Table Grid'
    _add_table_row(overview_table.rows[0], 'Algorithm Name', scenario_info.get('name', 'N/A'))
    _add_table_row(overview_table.rows[1], 'Version', scenario_info.get('version', 'N/A'))
    _add_table_row(overview_table.rows[2], 'PSDL Version', scenario_info.get('psdl_lang_version', 'N/A'))
    _add_table_row(overview_table.rows[3], 'Generation Date', datetime.now().strftime('%Y-%m-%d'))
    doc.add_paragraph()

    # Algorithm explanation
    doc.add_heading('How the Algorithm Works', level=2)
    if enriched_content and enriched_content.get('algorithm_explanation'):
        _add_narrative_paragraphs(doc, enriched_content['algorithm_explanation'])
    else:
        signals = scenario_info.get('signals', [])
        trends = scenario_info.get('trends', [])
        logic_rules = scenario_info.get('logic', [])
        doc.add_paragraph(
            f"This algorithm monitors {len(signals)} clinical data element(s), "
            f"computes {len(trends)} derived metric(s), and applies {len(logic_rules)} "
            f"detection rule(s) to identify the target clinical condition."
        )
    doc.add_paragraph()

    # Section 4: Data Elements
    doc.add_heading('4. Data Elements Required', level=1)

    # Narrative description
    if enriched_content and enriched_content.get('data_elements_narrative'):
        _add_narrative_paragraphs(doc, enriched_content['data_elements_narrative'])
        doc.add_paragraph()

    # Data elements table
    doc.add_heading('Data Element Specification', level=2)
    signals = scenario_info.get('signals', [])
    if signals:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Element'
        hdr[1].text = 'Clinical Reference'
        hdr[2].text = 'Unit'
        hdr[3].text = 'Description'
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True
            _set_cell_shading(cell, 'D9E2F3')

        for signal in signals:
            row = table.add_row().cells
            row[0].text = signal.get('name', '-')
            row[1].text = signal.get('ref', '-') or '-'
            row[2].text = signal.get('unit', '-') or '-'
            row[3].text = signal.get('description', '') or '-'
    else:
        _add_placeholder(doc, 'No data elements defined.')
    doc.add_paragraph()

    # Section 5: Algorithm Logic
    doc.add_heading('5. Algorithm Logic', level=1)

    # Trends
    trends = scenario_info.get('trends', [])
    if trends:
        doc.add_heading('Computed Metrics (Trends)', level=2)
        for trend in trends:
            p = doc.add_paragraph()
            p.add_run(f"{trend.get('name', 'Unknown')}: ").bold = True
            desc = trend.get('description') or f"Computed as {trend.get('expr', 'N/A')}"
            p.add_run(desc)
        doc.add_paragraph()

    # Logic rules
    logic_rules = scenario_info.get('logic', [])
    if logic_rules:
        doc.add_heading('Detection Rules', level=2)
        for rule in logic_rules:
            p = doc.add_paragraph()
            severity = rule.get('severity', 'info')
            severity_color = {
                'low': '28A745',
                'medium': 'FFC107',
                'high': 'FD7E14',
                'critical': 'DC3545'
            }.get(severity, '6C757D')

            name_run = p.add_run(f"{rule.get('name', 'Unknown')} ")
            name_run.bold = True

            severity_run = p.add_run(f"[{severity.upper()}]")
            severity_run.bold = True

            if rule.get('description'):
                doc.add_paragraph(f"Purpose: {rule.get('description')}")
            doc.add_paragraph(f"Condition: {rule.get('expr', 'N/A')}", style='Quote')
            doc.add_paragraph()
    doc.add_paragraph()

    # Section 6: Clinical Workflow
    doc.add_heading('6. Clinical Workflow Integration', level=1)
    if enriched_content and enriched_content.get('clinical_workflow'):
        _add_narrative_paragraphs(doc, enriched_content['clinical_workflow'])
    else:
        justification = governance_data.get('justification')
        if justification:
            doc.add_paragraph(justification)
        else:
            _add_placeholder(doc, 'Description of how this algorithm integrates into clinical workflows.')
    doc.add_paragraph()

    # Section 7: Safety Considerations
    doc.add_heading('7. Safety Considerations', level=1)
    if enriched_content and enriched_content.get('safety_considerations'):
        _add_narrative_paragraphs(doc, enriched_content['safety_considerations'])
    else:
        risk_assessment = governance_data.get('risk_assessment')
        if risk_assessment:
            doc.add_paragraph(risk_assessment)
        else:
            _add_placeholder(doc, 'Analysis of potential risks, false positives/negatives, and mitigation strategies.')
    doc.add_paragraph()

    # Section 8: Limitations
    doc.add_heading('8. Limitations', level=1)
    if enriched_content and enriched_content.get('limitations'):
        _add_narrative_paragraphs(doc, enriched_content['limitations'])
    else:
        _add_placeholder(doc, 'Known limitations and appropriate use cases for this algorithm.')
    doc.add_paragraph()

    # Section 9: Recommendation
    doc.add_heading('9. Recommendation', level=1)
    if enriched_content and enriched_content.get('recommendation'):
        _add_narrative_paragraphs(doc, enriched_content['recommendation'])
    else:
        _add_placeholder(doc, 'Final recommendation for IRB consideration.')
    doc.add_paragraph()

    # Appendix: Technical Specification
    doc.add_page_break()
    doc.add_heading('Appendix A: Technical Specification', level=1)

    doc.add_heading('A.1 Signal Definitions', level=2)
    if signals:
        for signal in signals:
            p = doc.add_paragraph(style='Quote')
            p.add_run(f"{signal.get('name')}: ").bold = True
            p.add_run(f"ref={signal.get('ref', 'N/A')}, unit={signal.get('unit', 'N/A')}")
            if signal.get('concept_id'):
                p.add_run(f", concept_id={signal.get('concept_id')}")

    doc.add_heading('A.2 Trend Expressions', level=2)
    if trends:
        for trend in trends:
            p = doc.add_paragraph(style='Quote')
            p.add_run(f"{trend.get('name')}: ").bold = True
            p.add_run(trend.get('expr', 'N/A'))

    doc.add_heading('A.3 Logic Expressions', level=2)
    if logic_rules:
        for rule in logic_rules:
            p = doc.add_paragraph(style='Quote')
            p.add_run(f"{rule.get('name')}: ").bold = True
            p.add_run(rule.get('expr', 'N/A'))

    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        "This document was generated by PSDL Inspector for IRB preparation purposes.\n"
        "Algorithm specification validated against psdl-lang standards.\n"
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _setup_styles(doc: Document) -> None:
    """Set up document styles."""
    # Normal style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Quote style for code/expressions
    if 'Quote' not in doc.styles:
        quote_style = doc.styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
    else:
        quote_style = doc.styles['Quote']
    quote_style.font.name = 'Consolas'
    quote_style.font.size = Pt(10)
    quote_style.paragraph_format.left_indent = Inches(0.5)


def _add_title_page(doc: Document, scenario_info: Dict[str, Any]) -> None:
    """Add title page to document."""
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('Clinical Algorithm Documentation')
    title_run.bold = True
    title_run.font.size = Pt(28)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('For Institutional Review Board (IRB) Review')
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()
    doc.add_paragraph()

    algo_name = doc.add_paragraph()
    algo_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = algo_name.add_run(scenario_info.get('name', 'Algorithm'))
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = RGBColor(0, 102, 204)

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver_run = version.add_run(f"Version {scenario_info.get('version', 'N/A')}")
    ver_run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date.add_run(datetime.now().strftime('%B %d, %Y'))
    date_run.font.size = Pt(12)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"Generated by PSDL Inspector\npsdl-lang v{scenario_info.get('psdl_lang_version', 'N/A')}")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()


def _add_narrative_paragraphs(doc: Document, text: str) -> None:
    """Add text as multiple paragraphs, splitting on double newlines."""
    paragraphs = text.strip().split('\n\n')
    for para in paragraphs:
        # Clean up single newlines within paragraph
        cleaned = ' '.join(para.split('\n'))
        doc.add_paragraph(cleaned)


def _add_placeholder(doc: Document, text: str) -> None:
    """Add placeholder text in italics."""
    p = doc.add_paragraph()
    run = p.add_run(f"[{text}]")
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)


def _add_table_row(row, label: str, value: str) -> None:
    """Helper to add a row to a table."""
    row.cells[0].text = label
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[1].text = value or 'N/A'


def _set_cell_shading(cell, color: str) -> None:
    """Set cell background color."""
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)
